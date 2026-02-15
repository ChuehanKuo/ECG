"""Convert the literature review markdown to a formatted DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# -- Helper functions --
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    if bold:
        run = p.add_run(text)
        run.bold = True
    else:
        # Handle inline bold markers **...**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title page
title = doc.add_heading('ICU Patient Outcome Prediction\nUsing Time-Series Clinical Data', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Comprehensive Literature Review & Project Planning Document')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run('Coverage: 2018 \u2013 2026')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# 1. LANDMARK STUDIES
# ============================================================
add_heading('1. Landmark Studies and Papers', 1)

add_heading('1.1 Foundational Benchmarks', 2)

add_paragraph('**Harutyunyan et al. (2019) \u2014 "Multitask Learning and Benchmarking with Clinical Time Series Data"**')
add_bullet('Authors: Hrayr Harutyunyan, Hrant Khachatrian, David C. Kale, Greg Ver Steeg, Aram Galstyan')
add_bullet('Publication: Scientific Data, Vol. 6, No. 1 (2019)')
add_bullet('Key Contribution: Proposed four standardized clinical prediction benchmarks using MIMIC-III: (1) in-hospital mortality, (2) physiologic decompensation, (3) length of stay, and (4) phenotype classification. Established a multi-task learning framework with channel-wise LSTMs and deep supervision.')
add_bullet('Results: LSTM-based models significantly outperformed linear baselines. Deep supervision improved decompensation and LOS predictions. Benchmark covers 40,000+ ICU stays.')

add_paragraph('**PhysioNet/CinC Challenge 2012 \u2014 ICU Mortality Prediction**')
add_bullet('Dataset: 12,000 adult patients from MIMIC-II; 5 general descriptors + 36 time-series variables from first 48 hours of ICU stay.')
add_bullet('Participation: 37 teams, ~200 submissions worldwide.')

add_paragraph('**PhysioNet/CinC Challenge 2019 \u2014 Early Prediction of Sepsis**')
add_bullet('Dataset: 60,000+ ICU patients with up to 40 clinical variables per hour; sourced from three hospital systems.')
add_bullet('Participation: 104 groups, 853 submissions.')
add_bullet('Key Finding: Models predict sepsis hours before clinical recognition, but cross-hospital generalizability remains a major challenge.')

add_heading('1.2 Google\'s Landmark EHR Study', 2)

add_paragraph('**Rajkomar et al. (2018) \u2014 "Scalable and Accurate Deep Learning with Electronic Health Records"**')
add_bullet('Authors: Alvin Rajkomar, Eyal Oren, Kai Chen, Andrew M. Dai et al. (Google, UCSF, Stanford, U. of Chicago)')
add_bullet('Data: 216,221 adult patients from two US academic medical centers; 46.8 billion data points in FHIR format.')
add_bullet('In-hospital mortality: AUROC 0.93\u20130.94')
add_bullet('30-day unplanned readmission: AUROC 0.75\u20130.76')
add_bullet('Key Innovation: Used raw, uncurated EHR data in FHIR format without site-specific harmonization.')

add_heading('1.3 Dynamic and Real-Time Prediction Studies', 2)

add_table(
    ['Model / Study', 'Year', 'Dataset', 'Observation Window', 'AUROC'],
    [
        ['TBAL (Time-Aware Bidirectional Attention LSTM)', '2025', 'MIMIC-IV', '12h\u20131d', '0.959'],
        ['TBAL', '2025', 'eICU-CRD', '12h\u20131d', '0.933'],
        ['RealMIP (Generative Imputation)', '2025', 'MIMIC-IV', 'Real-time', '0.968'],
        ['CRISP (Causal Transformer)', '2024', 'MIMIC-III', '48h', '0.948'],
        ['CRISP', '2024', 'MIMIC-IV', '48h', '0.917'],
        ['Real-Time Ensemble (DL + LightGBM)', '2024', 'Internal (Korea)', 'Real-time', '0.866'],
        ['Real-Time Ensemble', '2024', 'MIMIC (external)', 'Real-time', '0.746'],
        ['Dynamic Survival (uncurated)', '2020', 'MIMIC-III', '48h', '0.850'],
    ]
)

# ============================================================
# 2. NOTABLE MODELS AND SYSTEMS
# ============================================================
add_heading('2. Notable Models and Systems', 1)

add_heading('2.1 InSight by Dascena', 2)
add_bullet('FDA-cleared sepsis prediction algorithm.')
add_bullet('Mechanism: Gradient tree boosting using only 6 vital signs + age and GCS. Requires only 2 hours of preceding data.')
add_bullet('First sepsis screening system to exceed AUROC 0.90 with vital signs alone.')
add_bullet('Prospective trial showed earlier sepsis detection and reduced ICU length of stay.')

add_heading('2.2 APACHE / SOFA / NEWS vs. Machine Learning', 2)

add_table(
    ['Study', 'ML Model', 'ML AUROC', 'Traditional Score', 'Trad. AUROC'],
    [
        ['Sci. Reports (2022)', 'Random Forest', '0.945', 'APACHE II', '~0.85'],
        ['GBM (MIMIC-III)', 'GBM', '0.927', 'SAPS II', '0.809'],
        ['XMI-ICU (2025)', 'XGBoost', '0.920', 'APACHE IV', '0.737'],
        ['Real-Time Ensemble (2024)', 'DL + LGBM', '0.866', 'NEWS', 'Lower (p<0.001)'],
    ]
)

add_paragraph('**Key Findings:** ML models generally outperform traditional scores (AUROC >0.90 vs 0.70\u20130.85). Hybrid approaches using ML to enhance traditional scores show promise. Traditional scores remain valuable for simplicity and interpretability.')

add_heading('2.3 Work from Taiwan / NTUH', 2)
add_bullet('**NTUH CAD Prediction (Biomedicines, 2022):** Multi-center study using AI to predict coronary artery disease from 12-lead ECGs. 2,303 patients, 12,954 ECG records.')
add_bullet('**Taipei Medical University:** Topic model + gradient boosting for ICU mortality prediction using MIMIC. AUROC 0.928.')
add_bullet('**NTUH Emergency Triage (JMIR, 2024):** Interpretable deep learning for triage level (63%), hospitalization (82%), LOS (71%) accuracy.')

# ============================================================
# 3. STATE-OF-THE-ART ARCHITECTURES
# ============================================================
add_heading('3. State-of-the-Art Architectures', 1)

add_heading('3.1 Recurrent Neural Networks (LSTM / GRU)', 2)
add_paragraph('**GRU-D (Che et al., 2018):** Incorporates missing data patterns directly into GRU architecture via trainable decay mechanisms. Instead of imputing then predicting, GRU-D treats missingness as informative and decays hidden states toward empirical means between observations. One of the most cited papers in clinical time-series ML.')
add_paragraph('**VS-GRU:** Extends GRU-D by learning features of different variables separately, accounting for variable-specific missing rates (often >80%).')
add_paragraph('**Time-Aware LSTM Variants:** Extend the forget gate to logarithmic or cubic decay functions of time intervals.')

add_heading('3.2 Temporal Convolutional Networks (TCNs)', 2)
add_bullet('Causal convolutions with dilated kernels enable exponentially growing receptive fields.')
add_bullet('Highly parallelizable, lower memory than RNNs, efficient for real-time deployment.')
add_bullet('Attention-based TCN achieved AUROC 0.837 (48h mortality on MIMIC-III).')

add_heading('3.3 Transformer-Based Models', 2)

add_table(
    ['Model', 'Key Innovation', 'Venue', 'Performance'],
    [
        ['STraTS', 'Set-based Transformer over (time, var, val) triplets', 'ACM TKDD 2022', 'Outperformed GRU-D on MIMIC-III'],
        ['Raindrop', 'GNN message passing between sensors', 'ICLR 2022', '+11.4% F1 on PhysioNet P19'],
        ['mTAND', 'Multi-time continuous attention', 'ICLR 2021', 'Outperformed IP-Net, SeFT'],
        ['CRISP', 'Causal graph + Transformer', '2024', 'AUROC 0.948 (MIMIC-III)'],
        ['DuETT', 'Dual attention over time + event type', '2023', 'SOTA on MIMIC-IV'],
        ['ContiFormer', 'Continuous-time Transformer', 'NeurIPS 2023', 'Irregular time series'],
    ]
)

add_heading('3.4 Neural ODE-Based Models', 2)
add_bullet('**GRU-ODE-Bayes (NeurIPS 2019):** Continuous-time GRU with Bayesian updates for sporadic observations.')
add_bullet('**Latent ODE (NeurIPS 2019):** VAE with ODE-based dynamics in latent space. Naturally handles arbitrary time gaps.')
add_bullet('**Neural CDE (NeurIPS 2020):** Extended Neural ODEs for irregular time series. Best AUC when observational intensity used as feature.')

add_heading('3.5 Handling Irregular Sampling and Missing Data', 2)

add_table(
    ['Strategy', 'Representative Models', 'Mechanism'],
    [
        ['Decay mechanisms', 'GRU-D, Time-Aware LSTM', 'Modify RNN gates with time-decay functions'],
        ['Masking + time intervals', 'GRU-D, VS-GRU', 'Encode missingness as explicit model inputs'],
        ['Observation triplets', 'STraTS, SeFT', 'Treat each (time, variable, value) as a set element'],
        ['Graph message passing', 'Raindrop', 'Model inter-sensor dependencies with GNN'],
        ['Continuous-time ODEs', 'GRU-ODE-Bayes, Latent ODE', 'Define dynamics in continuous time'],
        ['Attention over time', 'mTAND, Transformers', 'Attend to arbitrary time points'],
        ['Generative imputation', 'RealMIP', 'Dynamic imputation via generative models'],
    ]
)

add_heading('3.6 Multi-Task Learning Approaches', 2)
add_bullet('**Harutyunyan et al. (2019):** Heterogeneous MTL with channel-wise LSTMs predicting mortality, decompensation, LOS, and phenotype simultaneously.')
add_bullet('**Frontiers in Medicine (2022):** RNN/GRU/LSTM with attention for simultaneous prediction \u2014 Mortality AUC 0.870, LOS AUC 0.765, Readmission AUC 0.635.')
add_bullet('**M3T-LM (2024):** Multi-modal multi-task learning model jointly predicting LOS and mortality.')

# ============================================================
# 4. KEY BENCHMARKS AND CHALLENGES
# ============================================================
add_heading('4. Key Benchmarks, Performance, and Challenges', 1)

add_heading('4.1 In-Hospital Mortality Prediction Summary', 2)

add_table(
    ['Model', 'Dataset', 'Window', 'AUROC'],
    [
        ['RealMIP', 'MIMIC-IV', 'Real-time', '0.968'],
        ['TBAL', 'MIMIC-IV', '12h\u20131d', '0.959'],
        ['CRISP', 'MIMIC-III', '48h', '0.948'],
        ['Random Forest', 'MIMIC-III', '24h', '0.945'],
        ['Google (Rajkomar)', 'Two US hospitals', 'Full stay', '0.93\u20130.94'],
        ['Gradient Boosting (Taiwan)', 'MIMIC', '\u2014', '0.928'],
        ['XMI-ICU', 'eICU', '6h', '0.920'],
        ['CRISP', 'MIMIC-IV', '48h', '0.917'],
        ['TGAM (Transformer)', 'MIMIC-III', '\u2014', '0.877'],
        ['RNN/GRU/LSTM + Attention', 'MIMIC-IV', '24h', '0.870'],
        ['Dynamic Survival', 'MIMIC-III', '48h', '0.850'],
        ['Attention-TCN', 'MIMIC-III', '48h', '0.837'],
    ]
)

add_heading('4.2 Most Important Predictive Features', 2)

add_paragraph('**Vital Signs (ranked by importance):**')
add_bullet('Heart rate and heart rate variability \u2014 among the strongest vital sign predictors')
add_bullet('Respiratory rate \u2014 consistently top-2 predictor')
add_bullet('Blood pressure (systolic and diastolic) \u2014 highest Chi-squared association with deterioration')
add_bullet('SpO2 (peripheral oxygen saturation)')
add_bullet('Temperature')
add_bullet('Glasgow Coma Scale (GCS) \u2014 top predictor across mortality, LOS, and readmission models')

add_paragraph('**Laboratory Values:**')
add_bullet('Blood urea nitrogen (BUN), Lactate, Anion gap, Platelet count, Bilirubin, Bicarbonate/pH, PTT')

add_paragraph('**Treatment Variables:**')
add_bullet('Norepinephrine (vasopressors), Invasive ventilation \u2014 including treatment variables adds up to +0.112 AUROC')

add_heading('4.3 Key Datasets', 2)

add_table(
    ['Dataset', 'Source', 'Size', 'Key Features'],
    [
        ['MIMIC-III', 'Beth Israel Deaconess, Boston', '~60,000 ICU stays', '17 vital signs, labs, notes, waveforms'],
        ['MIMIC-IV', 'Same (updated)', '~70,000+ ICU stays', 'Updated through 2022'],
        ['eICU-CRD', 'Philips multi-center', '200,000+ ICU stays', '208 hospitals across US'],
        ['PhysioNet P12', 'MIMIC-II subset', '12,000 patients', '36 time-series, 48h window'],
        ['PhysioNet P19', 'Multi-hospital', '38,803 patients', '34 sensors, sepsis labels'],
        ['AmsterdamUMCdb', 'Amsterdam UMC', '~20,000 admissions', 'European ICU data'],
    ]
)

add_heading('4.4 Major Challenges', 2)

add_paragraph('**1. Irregular Sampling and Missing Data**')
add_bullet('Clinical time series contain >80% missing values in some variables.')
add_bullet('Missingness is often "informative" \u2014 correlated with patient severity.')

add_paragraph('**2. Generalizability and External Validation**')
add_bullet('Performance consistently drops on external datasets (e.g., 0.878 \u2192 0.720 AUROC).')
add_bullet('~50% of published models lack external validation.')

add_paragraph('**3. Class Imbalance**')
add_bullet('ICU mortality rates are typically 10\u201315%, requiring focal loss or class weights.')

add_paragraph('**4. Clinical Deployment Barriers**')
add_bullet('Explainability, regulatory approval, alert fatigue, prospective validation.')

add_heading('4.5 Emerging Directions (2024\u20132026)', 2)
add_bullet('Causal inference integration (CRISP framework)')
add_bullet('Federated learning across institutions')
add_bullet('Foundation models for clinical time series')
add_bullet('Organ-specific outcome prediction (AKI, ventilator weaning, neurological)')
add_bullet('Multi-modal fusion (structured data + notes + imaging + waveforms)')
add_bullet('Latent SDE models for uncertainty quantification')

# ============================================================
# 5. PROPOSED PROJECT DESIGN
# ============================================================
doc.add_page_break()
add_heading('5. Proposed Project Design', 1)

add_heading('5.1 Pipeline Overview', 2)

add_paragraph('**Data Layer:**')
add_bullet('Raw ICU Data (vitals, labs, medications, demographics)')
add_bullet('Cohort Selection \u2192 Feature Extraction')
add_bullet('Temporal Binning (1h) \u2192 Outlier Clipping')
add_bullet('Normalization \u2192 Imputation \u2192 Missingness Masks')
add_bullet('Patient-Level Train/Val/Test Split (temporal)')

add_paragraph('**Model Layer:**')
add_bullet('Shared Temporal Encoder (GRU-D / Transformer / Mamba)')
add_bullet('Static Feature Embedding (demographics, admission info)')
add_bullet('Task-specific heads: Mortality, LOS, Anomaly/Decompensation')

add_paragraph('**Evaluation & Explainability Layer:**')
add_bullet('AUROC, AUPRC, Calibration (ECE/Brier Score)')
add_bullet('SHAP values, Attention visualization')
add_bullet('Lead-time analysis, Alert frequency analysis')

add_heading('5.2 Recommended Model Architecture', 2)
add_paragraph('**Primary:** Transformer-based encoder with time-aware attention (handles irregular sampling natively) + GRU-D as a strong baseline for comparison.')

add_paragraph('**Why Transformer over LSTM:**')
add_bullet('Parallelizable training')
add_bullet('Better at capturing long-range dependencies')
add_bullet('Natural attention-based interpretability')
add_bullet('SOTA results on MIMIC benchmarks (TBAL, CRISP, STraTS)')

add_heading('5.3 Feature Set', 2)

add_table(
    ['Category', 'Features', 'Encoding'],
    [
        ['Vitals', 'HR, SBP, DBP, MAP, SpO2, RR, Temp, GCS', 'Continuous (hourly mean + count)'],
        ['Labs', 'Creatinine, BUN, Lactate, WBC, Hgb, Platelets, pH, PaO2', 'Continuous'],
        ['Medications', 'Vasopressors, Sedatives, Antibiotics', 'Binary on/off + dose rate'],
        ['Demographics', 'Age, Sex', 'Static embedding'],
        ['Derived', 'Shock index (HR/SBP), P/F ratio, Fluid balance', 'Continuous'],
        ['Missingness', 'Binary mask per feature per timestep', 'Binary'],
    ]
)

add_heading('5.4 Training Strategy', 2)
add_bullet('**Multi-task loss:** Weighted sum of mortality (BCE), LOS (cross-entropy), decompensation (BCE)')
add_bullet('**Class imbalance:** Focal loss + class weights')
add_bullet('**Regularization:** Dropout, weight decay, early stopping on validation AUPRC')
add_bullet('**Validation:** Temporal split (train on earlier admissions, validate on later)')

# ============================================================
# 6. REFERENCES
# ============================================================
doc.add_page_break()
add_heading('6. References', 1)

refs = [
    'Harutyunyan et al. "Multitask Learning and Benchmarking with Clinical Time Series Data." Scientific Data 6, 96 (2019).',
    'Rajkomar et al. "Scalable and Accurate Deep Learning with Electronic Health Records." npj Digital Medicine 1, 18 (2018).',
    'Che et al. "Recurrent Neural Networks for Multivariate Time Series with Missing Values." Scientific Reports (2018).',
    'Tipirneni & Reddy. "STraTS: Self-Supervised Transformer for Sparse and Irregularly Sampled Clinical Time-Series." ACM TKDD 16(6), 2022.',
    'Zhang et al. "Raindrop: Graph-Guided Network for Irregularly Sampled Multivariate Time Series." ICLR 2022.',
    'Shukla & Marlin. "Multi-Time Attention Networks for Irregularly Sampled Time Series." ICLR 2021.',
    'De Brouwer et al. "GRU-ODE-Bayes: Continuous Modeling of Sporadically-Observed Time Series." NeurIPS 2019.',
    'Rubanova et al. "Latent ODEs for Irregularly-Sampled Time Series." NeurIPS 2019.',
    'Kidger et al. "Neural Controlled Differential Equations for Irregular Time Series." NeurIPS 2020.',
    '"TBAL: Dynamic Real-Time Risk Prediction Model for ICU Patients." JMIR (2025).',
    '"RealMIP: Unlocking the Potential of Real-Time ICU Mortality Prediction." npj Digital Medicine (2025).',
    '"CRISP: Causal Relationship Informed Superior Prediction." (2024).',
    'Dascena InSight. FDA-cleared sepsis prediction system.',
    '"Attention-Based TCN for ICU Mortality." BMC Anesthesiology (2022).',
    '"XMI-ICU: Explainable ML for ICU Mortality in MI Patients." Scientific Reports (2025).',
    '"NTUH AI-Enabled ECG for CAD Prediction." Biomedicines (2022).',
    '"YAIB: Yet Another ICU Benchmark." OpenReview (2023).',
    '"MIMIC-Sepsis Benchmark." arXiv (2025).',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.add_run(f'[{i}] ').bold = True
    p.add_run(ref)

# Save
doc.save('/home/user/ECG/ICU_Prediction_Literature_Review.docx')
print("DOCX saved successfully.")
