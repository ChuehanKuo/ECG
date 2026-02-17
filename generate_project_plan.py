"""Generate comprehensive NTUH ICU AI Project Plan DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# -- Helper functions --
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    if bold:
        run = p.add_run(text)
        run.bold = True
    else:
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def add_numbered(text, number):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{number}. ')
    run_num.bold = True
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('NTUH ICU Early Warning System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('AI-Driven Clinical Deterioration Prediction\nfor Preventive Intervention')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run('Comprehensive Project Plan')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run2.bold = True

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub3.add_run('National Taiwan University Hospital\n~20,000 Patients | Full-Stay Multi-Modal ICU Data')
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

sub4 = doc.add_paragraph()
sub4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = sub4.add_run('Target: Publication (Nature Digital Medicine / Lancet Digital Health) + Clinical Deployment')
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Why This Project Matters: The Preventive Care Gap',
    '3. Data Asset Assessment',
    '4. Primary Prediction Target: Clinical Deterioration',
    '5. Secondary Targets: Multi-Task Learning',
    '6. Model Architecture: Recommended Approach',
    '7. Complete Technical Pipeline',
    '8. Publication Strategy',
    '9. Clinical Deployment Roadmap',
    '10. TFDA Regulatory Pathway',
    '11. Ethical, IRB, and Data Governance',
    '12. Project Timeline',
    '13. Risk Mitigation',
    '14. Competitive Landscape & Differentiation',
    '15. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading('1. Executive Summary', 1)

add_paragraph('**The Opportunity:** NTUH possesses a uniquely valuable dataset \u2014 continuous, full-stay, multi-modal ICU data from ~20,000 patients including ECG waveforms, vital signs, blood pressure, laboratory values, and medication records. This is comparable in scale to AmsterdamUMCdb (~20,000 admissions) and represents one of the largest single-center Asian ICU datasets available for AI research.')

add_paragraph('**The Problem:** Despite advances in critical care, preventable deterioration in ICU patients remains a leading cause of adverse outcomes. Studies show that clinical deterioration events are often preceded by detectable physiological changes 6\u201324 hours before they become clinically obvious. Current early warning scores (APACHE, SOFA, NEWS) are static, updated infrequently, and achieve AUROCs of only 0.70\u20130.85.')

add_paragraph('**The Solution:** We propose building a real-time, multi-task AI early warning system that continuously monitors patient data streams and predicts clinical deterioration events \u2014 including hemodynamic instability, sepsis onset, respiratory failure, and cardiac arrest \u2014 hours before they occur, enabling preventive intervention.')

add_paragraph('**Key Differentiators:**')
add_bullet('**Full-stay temporal modeling:** Unlike most studies using only first 24\u201348 hours, we use the complete ICU trajectory, enabling dynamic risk re-assessment at every time step')
add_bullet('**ECG waveform integration:** Raw ECG data combined with structured clinical data (multi-modal fusion) \u2014 a frontier that few studies have explored at this scale')
add_bullet('**Asian population validation:** Critical gap in the literature; most models are trained on US/European data (MIMIC, eICU) and underperform on Asian populations')
add_bullet('**Dual purpose:** Designed simultaneously for high-impact publication AND real-world clinical deployment at NTUH')

add_paragraph('**Target Performance:** AUROC >0.90 for 6-hour deterioration prediction, with PPV >25% to minimize alert fatigue.')

doc.add_page_break()

# ============================================================
# 2. WHY THIS PROJECT MATTERS
# ============================================================
add_heading('2. Why This Project Matters: The Preventive Care Gap', 1)

add_heading('2.1 The Problem with Current Practice', 2)

add_paragraph('ICU clinicians currently rely on a combination of experience, periodic assessments, and traditional scoring systems to identify patients at risk of deterioration. This approach has fundamental limitations:')

add_bullet('**Delayed recognition:** Studies show 6\u201324 hours of detectable physiological decline precede ~80% of cardiac arrests, yet current systems detect deterioration only when it becomes clinically obvious')
add_bullet('**Static scores:** APACHE II/IV, SOFA, and NEWS are calculated at fixed intervals (typically once per shift), missing the dynamic evolution of patient state between assessments')
add_bullet('**Information overload:** ICU nurses monitor 200+ data points per patient per hour. Critical patterns are buried in noise')
add_bullet('**Alert fatigue:** Existing threshold-based alarms have >95% false positive rates, leading clinicians to ignore them')

add_heading('2.2 What Preventive Intervention Means', 2)

add_paragraph('The goal is not just prediction \u2014 it is **actionable early warning** that changes clinical behavior:')

add_table(
    ['Predicted Event', 'Lead Time Target', 'Preventive Action Enabled'],
    [
        ['Hemodynamic instability', '4\u20136 hours', 'Early fluid resuscitation, vasopressor initiation, ICU bed allocation'],
        ['Sepsis onset', '6\u201312 hours', 'Early blood cultures, empiric antibiotics (every hour of delay increases mortality 7.6%)'],
        ['Respiratory failure', '4\u20138 hours', 'Pre-emptive intubation preparation, non-invasive ventilation trial, positioning'],
        ['Cardiac arrest', '2\u20136 hours', 'Electrolyte correction, medication review, code team pre-alert'],
        ['Acute kidney injury', '12\u201348 hours', 'Nephrotoxin avoidance, fluid optimization, early nephrology consult'],
        ['Unexpected mortality', '24\u201348 hours', 'Goals-of-care discussion, family notification, palliative care consultation'],
    ]
)

add_heading('2.3 Evidence That AI Early Warning Systems Work', 2)

add_paragraph('Recent meta-analyses and prospective studies demonstrate that AI early warning systems can meaningfully improve outcomes:')

add_bullet('**2025 Meta-analysis (5 prospective studies):** AI early warning systems significantly reduce in-hospital mortality and 30-day mortality (pooled OR 0.72, 95% CI 0.58\u20130.89)')
add_bullet('**CMUH (Taiwan):** i.A.M.S. system reduced mortality by 25% and antibiotic costs by 30% across 12 hospitals')
add_bullet('**eCART (Yale New Haven Health):** Universal adoption with >90% nurse compliance; significant reduction in code blue events')
add_bullet('**Stanford Health Care:** AI deterioration model integrated into EHR reduced unexpected ICU transfers by 35%')
add_bullet('**Singapore BRAIN Platform:** SERA sepsis algorithm (AUROC 0.94 at 12h pre-onset) deployed across all public hospitals nationally since 2017')

doc.add_page_break()

# ============================================================
# 3. DATA ASSET ASSESSMENT
# ============================================================
add_heading('3. Data Asset Assessment', 1)

add_heading('3.1 Dataset Overview', 2)

add_table(
    ['Attribute', 'Details'],
    [
        ['Source', 'National Taiwan University Hospital (NTUH) ICU'],
        ['Patient Count', '~20,000 patients'],
        ['Coverage', 'Full ICU stay (admission to discharge/death)'],
        ['Temporal Resolution', 'Continuous (vital signs), periodic (labs), event-based (medications)'],
        ['Data Types', 'ECG waveforms, vital signs, blood pressure, laboratory values, medications, demographics'],
        ['Population', 'Taiwanese / East Asian adults'],
        ['Comparable Datasets', 'AmsterdamUMCdb (~20K), MIMIC-III (~60K stays), eICU-CRD (~200K stays)'],
    ]
)

add_heading('3.2 Expected Data Modalities', 2)

add_table(
    ['Category', 'Expected Variables', 'Typical Sampling Rate', 'Notes'],
    [
        ['ECG', '12-lead waveforms, derived HR, HRV, rhythm', 'Continuous (250\u2013500 Hz)', 'Unique advantage \u2014 most ICU datasets lack raw ECG'],
        ['Vital Signs', 'HR, SBP, DBP, MAP, SpO2, RR, Temperature', 'Every 1\u20135 minutes', 'Core features for deterioration prediction'],
        ['Laboratory', 'CBC, BMP, LFTs, Lactate, ABG, Coagulation', 'Every 4\u201324 hours', 'Irregular sampling \u2014 requires careful handling'],
        ['Medications', 'Vasopressors, sedatives, antibiotics, fluids', 'Event-based', 'Both indicator and predictor of state changes'],
        ['Demographics', 'Age, sex, admission diagnosis, comorbidities', 'Static', 'Context features'],
        ['Scores', 'APACHE, SOFA, GCS (if available)', 'Every 8\u201324 hours', 'Useful as comparison benchmarks'],
    ]
)

add_heading('3.3 Data Strengths and Advantages', 2)

add_bullet('**Full-stay data:** Unlike MIMIC benchmarks that typically use only the first 24\u201348 hours, full-stay data enables dynamic, evolving risk prediction throughout the ICU course')
add_bullet('**ECG waveforms:** Raw ECG provides sub-second cardiac dynamics that structured data cannot capture. ECG-derived features (HRV, QT interval, ST changes) are powerful predictors of cardiac events and sepsis')
add_bullet('**Medication records:** Treatment data (vasopressors, antibiotics) is both a predictor and confounder. Including it can add +0.112 AUROC (per recent studies) but requires careful causal handling')
add_bullet('**Single-center consistency:** Uniform data collection protocols, EHR systems, and clinical practices reduce confounding from inter-hospital variation')
add_bullet('**Asian population:** Addresses critical gap \u2014 all major ICU benchmarks (MIMIC, eICU, AmsterdamUMCdb) are Western populations. Physiological baselines differ (e.g., blood pressure norms, BMI distributions)')

add_heading('3.4 Data Challenges to Address', 2)

add_bullet('**Missing data:** ICU clinical time series typically have >80% missing values for some variables. Strategy: GRU-D decay mechanism + explicit missingness encoding')
add_bullet('**Irregular sampling:** Labs drawn every 4\u201324h vs vitals every 1\u20135min. Strategy: Multi-resolution temporal encoding')
add_bullet('**Label definition:** "Deterioration" must be precisely defined. Strategy: Composite endpoint (see Section 4)')
add_bullet('**Class imbalance:** Mortality ~10\u201315%, sepsis ~5\u201310%. Strategy: Focal loss + AUPRC as primary metric')
add_bullet('**Temporal confounding:** Treatments affect outcomes. Strategy: Inverse probability weighting or causal transformer (CRISP)')

doc.add_page_break()

# ============================================================
# 4. PRIMARY PREDICTION TARGET
# ============================================================
add_heading('4. Primary Prediction Target: Clinical Deterioration', 1)

add_heading('4.1 Why "Clinical Deterioration" Over Single Endpoints', 2)

add_paragraph('We recommend a **composite clinical deterioration** endpoint rather than predicting mortality or sepsis alone. Here is why:')

add_table(
    ['Approach', 'Pros', 'Cons'],
    [
        ['Mortality only', 'Clean label, easy to validate', 'Too late for prevention; low event rate (~10%)'],
        ['Sepsis only', 'High clinical impact', 'Time-Zero problem (Epic Sepsis Model failure); definition controversy'],
        ['Single organ failure', 'Specific, actionable', 'Narrow scope; misses other deterioration patterns'],
        ['Composite deterioration (RECOMMENDED)', 'Captures all preventable events; higher event rate; avoids Time-Zero trap', 'Requires careful endpoint definition; heterogeneous outcomes'],
    ]
)

add_heading('4.2 Recommended Composite Endpoint Definition', 2)

add_paragraph('**Primary outcome: Acute clinical deterioration** defined as the FIRST occurrence of any of the following within the prediction window:')

add_numbered('**New vasopressor initiation** (norepinephrine, vasopressin, epinephrine, dopamine, phenylephrine) \u2014 indicates hemodynamic instability', 1)
add_numbered('**New mechanical ventilation** (intubation) \u2014 indicates respiratory failure', 2)
add_numbered('**Cardiac arrest / Code Blue** \u2014 most severe deterioration', 3)
add_numbered('**Rapid Response Team activation** (if recorded) \u2014 clinician-identified deterioration', 4)
add_numbered('**In-ICU death** \u2014 ultimate negative outcome', 5)
add_numbered('**Acute kidney injury** (KDIGO Stage 2+: creatinine >2x baseline or UO <0.5 mL/kg/h for 12h)', 6)

add_paragraph('**Prediction horizon:** 6-hour lookahead (primary), with secondary analyses at 4h, 8h, 12h, and 24h.')

add_paragraph('**Why this endpoint works:**')
add_bullet('Avoids the "Time-Zero" problem that plagued Epic Sepsis Model (we predict new interventions, not recognition)')
add_bullet('Higher event rate than mortality alone (~20\u201330% of patients experience at least one), improving model training')
add_bullet('Every component is directly actionable \u2014 clinicians can take specific preventive measures')
add_bullet('Aligns with eCART (University of Chicago) approach, which achieved real-world clinical success')

doc.add_page_break()

# ============================================================
# 5. SECONDARY TARGETS: MULTI-TASK LEARNING
# ============================================================
add_heading('5. Secondary Targets: Multi-Task Learning', 1)

add_paragraph('Multi-task learning (MTL) improves both performance and clinical utility by sharing representations across related tasks. We recommend jointly predicting:')

add_table(
    ['Task', 'Label Definition', 'Loss Function', 'Clinical Utility'],
    [
        ['Deterioration (primary)', 'Composite endpoint within 6h', 'Focal BCE', 'Trigger early intervention'],
        ['ICU Mortality', 'Death during ICU stay', 'BCE', 'Goals-of-care discussions, resource allocation'],
        ['Remaining LOS', 'Hours until ICU discharge', 'Huber loss (regression)', 'Bed management, discharge planning'],
        ['Sepsis onset', 'Sepsis-3 criteria within 6h', 'Focal BCE', 'Early antibiotic initiation'],
        ['AKI progression', 'KDIGO stage progression within 48h', 'Ordinal CE', 'Nephrotoxin avoidance, renal consult'],
    ]
)

add_paragraph('**Multi-task training strategy:** Uncertainty-weighted loss (Kendall et al., 2018) to automatically balance task losses based on homoscedastic uncertainty:')
add_paragraph('L_total = (1/2\u03c3\u00b2_1) * L_deterioration + (1/2\u03c3\u00b2_2) * L_mortality + (1/2\u03c3\u00b2_3) * L_LOS + ... + log(\u03c3_1\u03c3_2\u03c3_3...)')

add_paragraph('**Expected benefits of MTL:**')
add_bullet('Shared temporal encoder learns more robust representations')
add_bullet('Regularization effect reduces overfitting on rare events (mortality)')
add_bullet('Single model serves multiple clinical needs, simplifying deployment')
add_bullet('Literature shows MTL improves mortality AUC by +0.02\u20130.05 over single-task models')

doc.add_page_break()

# ============================================================
# 6. MODEL ARCHITECTURE
# ============================================================
add_heading('6. Model Architecture: Recommended Approach', 1)

add_heading('6.1 Architecture Overview', 2)

add_paragraph('We recommend a **three-tier architecture** that processes data at multiple resolutions:')

add_paragraph('**Tier 1: ECG Waveform Encoder (High-Frequency)**')
add_bullet('Input: Raw ECG waveforms (250\u2013500 Hz)')
add_bullet('Architecture: 1D ResNet or Temporal Convolutional Network (TCN)')
add_bullet('Output: Per-minute ECG embedding vectors (capturing rhythm, morphology, HRV)')
add_bullet('Rationale: Compresses high-frequency waveform into manageable representations')

add_paragraph('**Tier 2: Clinical Time-Series Encoder (Medium-Frequency)**')
add_bullet('Input: Hourly-binned vitals + labs + medication indicators + ECG embeddings from Tier 1 + missingness masks + time-since-last-observation')
add_bullet('Architecture: Transformer with continuous time-aware positional encoding (STraTS-style) OR GRU-D as baseline')
add_bullet('Output: Per-hour patient state representation')
add_bullet('Rationale: Handles irregular sampling natively; captures long-range dependencies across the full stay')

add_paragraph('**Tier 3: Task-Specific Prediction Heads**')
add_bullet('Input: Patient state representation from Tier 2 + static features (age, sex, admission diagnosis)')
add_bullet('Architecture: MLP heads with dropout, one per task')
add_bullet('Output: Per-hour predictions for each task (deterioration probability, mortality risk, expected remaining LOS)')

add_heading('6.2 Detailed Architecture: Clinical Time-Series Encoder', 2)

add_paragraph('**Option A (Primary): Time-Aware Transformer**')
add_bullet('Multi-head self-attention with learnable continuous-time positional encoding')
add_bullet('Each observation is a triplet: (timestamp, variable_id, value) \u2014 following STraTS design')
add_bullet('Variable-specific embedding layer captures different dynamics per feature')
add_bullet('Relative time encoding via sinusoidal functions of inter-observation intervals')
add_bullet('Expected AUROC: 0.91\u20130.95 (based on CRISP, STraTS, TBAL benchmarks)')

add_paragraph('**Option B (Baseline): GRU-D**')
add_bullet('GRU with trainable exponential decay for missing values')
add_bullet('Missingness masks and time-since-last-observation as explicit inputs')
add_bullet('Well-established, highly cited, excellent baseline for comparison')
add_bullet('Expected AUROC: 0.85\u20130.90')

add_paragraph('**Option C (Advanced): Mamba / State Space Model**')
add_bullet('Linear-time sequence modeling (O(n) vs O(n\u00b2) for Transformer)')
add_bullet('Emerging architecture with strong results on long sequences')
add_bullet('Would be highly novel for ICU prediction \u2014 potential for ICML/NeurIPS paper')
add_bullet('Risk: Less established in clinical domain; may require more engineering')

add_heading('6.3 ECG-Specific Processing', 2)

add_paragraph('The ECG waveform component is a key differentiator. Recommended approach:')

add_bullet('**Pre-processing:** Bandpass filter (0.5\u201340 Hz), baseline wander removal, R-peak detection')
add_bullet('**Feature extraction (classical):** Heart rate variability (SDNN, RMSSD, pNN50), QT interval, ST segment deviation, QRS duration')
add_bullet('**Feature extraction (learned):** 1D ResNet-18 operating on 10-second ECG segments, producing 128-dim embeddings')
add_bullet('**Fusion strategy:** Concatenate ECG embeddings with structured clinical features at each hourly timestep')
add_bullet('**Clinical value:** ECG-derived features can detect subclinical myocardial ischemia, electrolyte imbalances (hyperkalemia), and autonomic dysfunction hours before they manifest in vital signs')

add_heading('6.4 Handling Irregular Sampling', 2)

add_table(
    ['Data Type', 'Sampling Rate', 'Strategy'],
    [
        ['ECG waveforms', '250\u2013500 Hz', 'TCN/ResNet encoder \u2192 per-minute embeddings'],
        ['Vital signs', '1\u20135 min', 'Aggregate to hourly: mean, std, min, max, count'],
        ['Laboratory values', '4\u201324 hours', 'Forward-fill + decay mask + time-since-last-observation'],
        ['Medications', 'Event-based', 'Binary on/off per hour + cumulative dose'],
        ['Missing values', 'Variable', 'Explicit missingness mask as input feature (missingness is informative)'],
    ]
)

add_heading('6.5 Interpretability', 2)

add_paragraph('For clinical deployment, the model must explain its predictions:')

add_bullet('**Attention weights:** Transformer attention maps show which time steps and variables drove the prediction')
add_bullet('**SHAP values:** TreeSHAP for gradient boosting baseline; DeepSHAP for neural models')
add_bullet('**Contributing factors display:** Top 3\u20135 features and their recent trends shown with each alert (e.g., "Rising lactate: 1.2 \u2192 2.8 mmol/L over 4h")')
add_bullet('**Temporal importance:** Which time window contributed most (recent vs. early stay)')

doc.add_page_break()

# ============================================================
# 7. COMPLETE TECHNICAL PIPELINE
# ============================================================
add_heading('7. Complete Technical Pipeline', 1)

add_heading('7.1 Data Preprocessing', 2)

add_numbered('**Cohort selection:** Adults (>=18), ICU stay >=4 hours, exclude readmissions within 24h (or treat as separate encounters)', 1)
add_numbered('**Variable harmonization:** Map NTUH-specific codes to standardized names; unit conversions', 2)
add_numbered('**Outlier handling:** Physiologically implausible values (e.g., HR <10 or >300, SBP <20 or >300) \u2192 clip or remove', 3)
add_numbered('**Temporal binning:** Aggregate to 1-hour windows (vitals: mean/std/min/max/count; labs: last value; meds: binary + dose)', 4)
add_numbered('**Normalization:** Per-variable z-score normalization using training set statistics', 5)
add_numbered('**Missingness encoding:** Binary mask matrix M (1=observed, 0=missing) + time-since-last-observation matrix \u0394t', 6)
add_numbered('**Label extraction:** For each hour t, compute whether deterioration endpoint occurs in [t+1, t+6]', 7)

add_heading('7.2 Data Splitting Strategy', 2)

add_paragraph('**Critical: Use temporal splitting, not random splitting.**')
add_bullet('**Training set (60%):** Patients admitted in the earlier time period')
add_bullet('**Validation set (15%):** Patients admitted in the middle time period')
add_bullet('**Test set (25%):** Patients admitted in the most recent time period')
add_bullet('**Rationale:** Temporal splitting simulates real-world deployment where the model must predict on future patients. Random splitting causes data leakage and inflates metrics.')

add_paragraph('Additional validation strategies:')
add_bullet('5-fold temporal cross-validation for robust performance estimates')
add_bullet('Subgroup analysis: by age group, sex, primary diagnosis, ICU type (medical vs. surgical)')
add_bullet('Calibration analysis: reliability diagrams, expected calibration error (ECE), Brier score')

add_heading('7.3 Training Configuration', 2)

add_table(
    ['Hyperparameter', 'Recommended Value', 'Notes'],
    [
        ['Learning rate', '1e-4 (Transformer) / 1e-3 (GRU-D)', 'With cosine annealing + warm-up'],
        ['Batch size', '64\u2013128 patients', 'Gradient accumulation if GPU memory limited'],
        ['Sequence length', 'Full stay (padded to max)', 'With attention mask for padding'],
        ['Hidden dimension', '128\u2013256', 'Larger for Transformer, smaller for GRU-D'],
        ['Number of heads', '4\u20138', 'Transformer only'],
        ['Number of layers', '4\u20136', 'Transformer; 2\u20133 for GRU-D'],
        ['Dropout', '0.2\u20130.3', 'On embeddings and between layers'],
        ['Loss function', 'Focal loss (\u03b1=0.25, \u03b3=2)', 'For class imbalance'],
        ['Optimizer', 'AdamW (weight decay 0.01)', 'Standard for Transformers'],
        ['Early stopping', 'Patience 10 on validation AUPRC', 'AUPRC better than AUROC for imbalanced data'],
        ['Max epochs', '100', 'Rarely needed with early stopping'],
    ]
)

add_heading('7.4 Evaluation Metrics', 2)

add_paragraph('**Primary metrics (must report all):**')
add_bullet('**AUROC:** Discrimination ability (threshold-independent)')
add_bullet('**AUPRC:** Critical for imbalanced outcomes \u2014 more informative than AUROC when positive rate <15%')
add_bullet('**Sensitivity at fixed specificity:** Sensitivity @ 95% specificity (clinical operating point)')
add_bullet('**PPV at clinical threshold:** Must achieve PPV >20\u201325% to be clinically useful (avoid alert fatigue)')

add_paragraph('**Secondary metrics:**')
add_bullet('**Calibration:** Expected Calibration Error (ECE), Brier score, reliability diagrams')
add_bullet('**Lead time analysis:** How far in advance does the model correctly predict deterioration?')
add_bullet('**Alert burden:** Number of alerts per patient per day at clinical operating threshold')
add_bullet('**Net reclassification improvement (NRI):** Improvement over APACHE/SOFA/NEWS')
add_bullet('**Subgroup performance:** Stratified by age, sex, diagnosis, ICU type')

add_heading('7.5 Baseline Comparisons', 2)

add_table(
    ['Baseline', 'Description', 'Expected AUROC'],
    [
        ['APACHE II/IV', 'Standard ICU severity score', '0.70\u20130.80'],
        ['SOFA score', 'Organ failure assessment', '0.70\u20130.80'],
        ['NEWS / qSOFA', 'Bedside screening score', '0.65\u20130.75'],
        ['Logistic Regression', 'Last-value features + demographics', '0.75\u20130.82'],
        ['XGBoost', 'Gradient boosting on engineered features', '0.85\u20130.90'],
        ['GRU-D', 'Decay-aware RNN baseline', '0.85\u20130.90'],
        ['Proposed Transformer', 'Full model', '0.91\u20130.95'],
    ]
)

doc.add_page_break()

# ============================================================
# 8. PUBLICATION STRATEGY
# ============================================================
add_heading('8. Publication Strategy', 1)

add_heading('8.1 Target Journals (Ranked by Impact and Fit)', 2)

add_table(
    ['Journal', 'Impact Factor', 'Fit', 'Review Time', 'Strategy'],
    [
        ['Nature Medicine', '82.9', 'If clinical trial included', '2\u20134 months', 'Highest impact; requires prospective validation'],
        ['Lancet Digital Health', '36.6', 'Excellent fit', '6\u20138 weeks', 'Primary target \u2014 perfect scope match'],
        ['npj Digital Medicine', '15.2', 'Excellent fit', '4\u20136 weeks', 'Strong backup; published Google EHR study'],
        ['Nature Communications', '16.6', 'Good for methods', '2\u20134 months', 'If novel architecture (Mamba/SSM)'],
        ['JAMA Network Open', '13.8', 'Clinical focus', '4\u20138 weeks', 'If strong clinical validation component'],
        ['Critical Care Medicine', '9.3', 'ICU audience', '6\u20138 weeks', 'Reaches ICU clinicians directly'],
    ]
)

add_heading('8.2 What Makes This Publishable in Top Journals', 2)

add_paragraph('Top-tier publication requires **novelty + rigor + clinical significance**. Here is how we achieve all three:')

add_paragraph('**Novelty:**')
add_bullet('First large-scale (20K patient) Asian ICU deterioration prediction study with full-stay multi-modal data')
add_bullet('ECG waveform + structured data fusion for ICU deterioration \u2014 underexplored in the literature')
add_bullet('Addresses the critical generalizability gap: Asian population-specific model development')
add_bullet('If using Mamba/SSM architecture: first application to clinical deterioration prediction')

add_paragraph('**Rigor:**')
add_bullet('Temporal train/test split (not random \u2014 many published papers make this error)')
add_bullet('Comprehensive baselines: APACHE, SOFA, NEWS, logistic regression, XGBoost, GRU-D, Transformer')
add_bullet('Calibration analysis (ECE, reliability diagrams) \u2014 often missing from ICU prediction papers')
add_bullet('Subgroup fairness analysis by age, sex, diagnosis')
add_bullet('TRIPOD+AI reporting guidelines compliance')

add_paragraph('**Clinical Significance:**')
add_bullet('Composite deterioration endpoint that maps to specific preventive actions')
add_bullet('Lead-time analysis showing the window for intervention')
add_bullet('Alert burden analysis at clinical operating thresholds')
add_bullet('Interpretability analysis showing clinically meaningful contributing factors')

add_heading('8.3 Recommended Paper Structure', 2)

add_paragraph('**Title (draft):** "Multi-Modal Deep Learning for Real-Time Clinical Deterioration Prediction in the ICU: Development and Temporal Validation Using 20,000 Full-Stay Admissions from a Taiwanese Medical Center"')

add_paragraph('**Suggested sections:**')
add_numbered('Introduction: Gap in Asian ICU prediction; limitations of static scores; need for real-time multi-modal systems', 1)
add_numbered('Methods: Data description, preprocessing, architecture, training, evaluation (TRIPOD+AI compliant)', 2)
add_numbered('Results: Main performance, temporal validation, subgroup analysis, calibration, lead-time, alert burden', 3)
add_numbered('Discussion: Comparison to MIMIC/eICU-trained models, clinical implications, limitations', 4)
add_numbered('Supplementary: Full architecture details, hyperparameter sensitivity, additional subgroup analyses', 5)

add_heading('8.4 Multi-Paper Strategy', 2)

add_paragraph('This dataset and system can support 3\u20135 publications:')

add_table(
    ['Paper', 'Focus', 'Target Journal', 'Timeline'],
    [
        ['Paper 1 (Primary)', 'Full system: multi-modal deterioration prediction', 'Lancet Digital Health / npj Digital Medicine', 'Month 6\u20138'],
        ['Paper 2', 'ECG-specific: waveform features for subclinical deterioration', 'European Heart Journal - Digital Health', 'Month 8\u201310'],
        ['Paper 3', 'Cross-population: NTUH vs MIMIC generalizability analysis', 'Critical Care Medicine', 'Month 10\u201312'],
        ['Paper 4', 'Clinical deployment: prospective pilot study results', 'JAMA Network Open / Nature Medicine', 'Month 18\u201324'],
        ['Paper 5', 'Methods: novel architecture (Mamba/SSM for clinical time series)', 'NeurIPS / ICML / AAAI', 'Month 8\u201312'],
    ]
)

doc.add_page_break()

# ============================================================
# 9. CLINICAL DEPLOYMENT ROADMAP
# ============================================================
add_heading('9. Clinical Deployment Roadmap', 1)

add_heading('9.1 NTUH Integration Architecture', 2)

add_paragraph('NTUH uses a Service-Oriented Architecture (SOA) with HL7 messaging middleware. The deployment architecture should leverage this existing infrastructure:')

add_paragraph('**Data Ingestion Layer:**')
add_bullet('Subscribe to HL7 ADT (admission/discharge/transfer), ORU (observation results), and RAS (pharmacy) messages via NTUH middleware')
add_bullet('ECG waveform streaming via DICOM or vendor-specific protocol from bedside monitors')
add_bullet('Real-time data aggregation into hourly feature vectors')

add_paragraph('**Inference Layer:**')
add_bullet('Model serving via NVIDIA Triton Inference Server (NTUH has NVIDIA AI supercomputers)')
add_bullet('Inference latency target: <2 seconds per patient per prediction cycle')
add_bullet('Prediction refresh: every 15 minutes (following Stanford model) or every hour')
add_bullet('GPU inference with CPU fallback for redundancy')

add_paragraph('**Alert Delivery Layer:**')
add_bullet('EHR-integrated alert: pop-up in NTUH HIS with risk score, contributing factors, and trend visualization')
add_bullet('Nursing station dashboard: real-time ward overview showing all patients\' risk levels')
add_bullet('Mobile push notification: for attending physicians via hospital secure messaging app')
add_bullet('Escalation pathway: low risk (dashboard only) \u2192 medium risk (dashboard + EHR flag) \u2192 high risk (EHR + mobile alert + charge nurse notification)')

add_heading('9.2 Alert Design to Minimize Fatigue', 2)

add_paragraph('Alert fatigue is the primary reason clinical AI systems fail. Our design principles:')

add_bullet('**Tiered alerting:** Three risk levels (Low/Medium/High) with different delivery channels')
add_bullet('**PPV threshold:** Only fire alerts when PPV >25% at the chosen operating point')
add_bullet('**Temporal suppression:** Minimum 2-hour gap between alerts for the same patient (unless risk score increases significantly)')
add_bullet('**Context-aware suppression:** Suppress alerts for patients already receiving the predicted intervention (e.g., don\'t alert "vasopressor needed" if already on vasopressors)')
add_bullet('**Contributing factors:** Every alert displays the top 3\u20135 features driving the prediction with temporal trends')
add_bullet('**Target: <3 high-priority alerts per nurse per 12-hour shift**')

add_heading('9.3 Deployment Phases', 2)

add_paragraph('**Phase 1: Silent Mode (Months 12\u201315)**')
add_bullet('Model runs in background, predictions logged but NOT shown to clinicians')
add_bullet('Prospective data collection for performance validation')
add_bullet('Monitor model calibration and drift in real-time')
add_bullet('Compare predictions to actual outcomes retrospectively')

add_paragraph('**Phase 2: Dashboard Mode (Months 15\u201318)**')
add_bullet('Risk scores displayed on nursing station dashboard')
add_bullet('No active alerts \u2014 clinicians can check proactively')
add_bullet('Collect usage data: how often do clinicians check? Do they change behavior?')
add_bullet('Qualitative feedback from nurses and physicians')

add_paragraph('**Phase 3: Active Alert Mode (Months 18\u201324)**')
add_bullet('Full alert system activated (tiered, with fatigue mitigation)')
add_bullet('Prospective clinical outcome comparison: pre- vs post-deployment')
add_bullet('Primary endpoint: reduction in unexpected deterioration events')
add_bullet('Secondary endpoints: time-to-intervention, ICU mortality, LOS')

add_paragraph('**Phase 4: Multi-Ward Expansion (Months 24+)**')
add_bullet('Extend to NTUH step-down units, emergency department')
add_bullet('Extend to NTUH branch hospitals')
add_bullet('Continuous learning: periodic model retraining with new data')

add_heading('9.4 Model Monitoring in Production', 2)

add_bullet('**Performance monitoring:** Weekly AUROC/AUPRC on recent predictions vs outcomes')
add_bullet('**Data drift detection:** Population Stability Index (PSI) and Kolmogorov\u2013Smirnov tests on input feature distributions')
add_bullet('**Calibration monitoring:** Weekly calibration curves \u2014 alert if ECE >0.05')
add_bullet('**Automated retraining trigger:** If AUROC drops >0.03 or PSI >0.2 on any critical feature')
add_bullet('**Human-in-the-loop:** Quarterly review by clinical champions (ICU attending + data science team)')

doc.add_page_break()

# ============================================================
# 10. TFDA REGULATORY PATHWAY
# ============================================================
add_heading('10. TFDA Regulatory Pathway', 1)

add_heading('10.1 Classification', 2)

add_paragraph('An AI clinical deterioration prediction system that **assists** clinicians (does not make autonomous decisions) would likely be classified as:')

add_bullet('**TFDA Class II medical device** (moderate risk)')
add_bullet('Category: Software as a Medical Device (SaMD)')
add_bullet('Comparable: EverFortune.AI sepsis prediction (TFDA-approved), CMUH i.A.M.S. system')

add_heading('10.2 Regulatory Timeline', 2)

add_table(
    ['Phase', 'Duration', 'Key Activities'],
    [
        ['Pre-submission consultation', '2\u20134 weeks', 'Meet with TFDA to confirm classification and requirements'],
        ['Quality System Documentation (QSD)', '6\u201310 months (concurrent with development)', 'ISO 13485 compliance, design controls, risk management (ISO 14971)'],
        ['Clinical validation', '6\u201312 months', 'Prospective silent-mode study at NTUH'],
        ['Product registration review', '4\u20137 months (Class II)', 'Technical file, clinical evidence, local validation data'],
        ['Total estimated', '12\u201824 months from start', 'Concurrent activities reduce wall-clock time'],
    ]
)

add_heading('10.3 Key TFDA Requirements (August 2025 Guidelines)', 2)

add_bullet('**Local population validation:** TFDA explicitly requires validation on Taiwanese population data (we have this)')
add_bullet('**Clinical significance documentation:** Must demonstrate that AI outputs have clear clinical meaning and utility')
add_bullet('**Lifecycle quality management:** ISO 13485-compliant QMS throughout development and post-market')
add_bullet('**Independent performance assessment:** Third-party or held-out temporal validation')
add_bullet('**Software lifecycle documentation:** IEC 62304 compliance for software development')

add_paragraph('**Advantage: NTUH has successfully navigated TFDA approval before** (PANCREASaver for pancreatic cancer screening). The institutional knowledge and regulatory relationships are in place.')

add_heading('10.4 Fast-Track Opportunities', 2)

add_bullet('**Government-funded projects:** If funded by NTU, MOST, or MOHW \u2192 eligible for priority review')
add_bullet('**Breakthrough designation:** If the system demonstrates substantial improvement over existing practice')
add_bullet('**Reciprocal recognition:** If simultaneously pursuing US FDA or EU CE, TFDA offers expedited review')

doc.add_page_break()

# ============================================================
# 11. ETHICAL, IRB, AND DATA GOVERNANCE
# ============================================================
add_heading('11. Ethical, IRB, and Data Governance', 1)

add_heading('11.1 IRB Approval Strategy', 2)

add_paragraph('NTUH has four Research Ethics Committees (REC A\u2013D), each meeting monthly with 19\u201320 members.')

add_paragraph('**Phase 1 \u2014 Retrospective Model Development:**')
add_bullet('Submit for **expedited review or exemption** (retrospective, de-identified data)')
add_bullet('Justification: secondary use of existing clinical data, minimal risk, no patient contact')
add_bullet('Expected timeline: 2\u20134 weeks for expedited; 4\u20136 weeks for full board if required')

add_paragraph('**Phase 2 \u2014 Prospective Validation (Silent Mode):**')
add_bullet('Submit for **full board review** (prospective study, even though model is silent)')
add_bullet('Requires: study protocol, data management plan, informed consent waiver justification')
add_bullet('Argument for consent waiver: model runs silently, no change to clinical care, minimal risk')

add_paragraph('**Phase 3 \u2014 Active Deployment Study:**')
add_bullet('**Full board review** with detailed protocol')
add_bullet('Requires: informed consent process (or waiver with strong justification), DSMB consideration')
add_bullet('May need cluster-randomized or stepped-wedge design for rigorous evaluation')

add_heading('11.2 Data Protection Compliance', 2)

add_paragraph('**Taiwan Personal Data Protection Act (PDPA):**')
add_bullet('Medical data is "sensitive personal data" requiring written informed consent for collection/processing')
add_bullet('**Key exemption:** 2022 Constitutional Court ruling upheld that de-identified health data can be used for secondary research')
add_bullet('De-identification requirements: remove direct identifiers (name, ID, MRN); k-anonymity for quasi-identifiers (dates, ages)')

add_paragraph('**2025 PDPA Amendments (new requirements):**')
add_bullet('Mandatory data breach notification')
add_bullet('Data Protection Officers required for government agencies')
add_bullet('Independent Personal Data Protection Commission established')

add_paragraph('**Taiwan AI Basic Law (December 2025):**')
add_bullet('Seven principles: transparency, privacy, autonomy, fairness, cybersecurity, sustainability, accountability')
add_bullet('Framework statute; detailed healthcare AI implementation regulations forthcoming')
add_bullet('Currently no specific informed consent requirement for AI-assisted clinical decision support beyond standard PDPA')

add_heading('11.3 Recommended Data Governance Framework', 2)

add_bullet('All data stored on NTUH secure servers (no cloud; NTUH has on-premise NVIDIA infrastructure)')
add_bullet('De-identification pipeline: remove names, IDs, MRNs; shift dates by random offset per patient')
add_bullet('Access control: role-based access with audit logging')
add_bullet('Data use agreement: formal agreement between research team and NTUH data governance office')
add_bullet('Model cards: document model purpose, training data, performance, limitations, and intended use')

doc.add_page_break()

# ============================================================
# 12. PROJECT TIMELINE
# ============================================================
add_heading('12. Project Timeline', 1)

add_table(
    ['Phase', 'Months', 'Key Deliverables'],
    [
        ['Phase 0: Setup & IRB', '0\u20132', 'IRB approval, data access agreement, compute setup, team assembly'],
        ['Phase 1: Data Engineering', '1\u20134', 'Data extraction, cleaning, cohort definition, feature engineering pipeline, EDA report'],
        ['Phase 2: Baseline Models', '3\u20135', 'APACHE/SOFA/NEWS baselines, logistic regression, XGBoost \u2014 establish performance floor'],
        ['Phase 3: Deep Learning Models', '4\u20137', 'GRU-D baseline, Transformer model, ECG encoder, multi-task training, hyperparameter tuning'],
        ['Phase 4: Evaluation & Analysis', '6\u20138', 'Temporal validation, subgroup analysis, calibration, lead-time, interpretability'],
        ['Phase 5: Paper Writing', '7\u20139', 'Draft primary paper, supplementary materials, submit to target journal'],
        ['Phase 6: Deployment Prep', '8\u201312', 'TFDA pre-submission, HL7 integration, inference pipeline, alert system design'],
        ['Phase 7: Silent Deployment', '12\u201315', 'Real-time inference, prospective validation, drift monitoring'],
        ['Phase 8: Active Deployment', '15\u201324', 'Dashboard \u2192 active alerts, clinical outcome evaluation, deployment paper'],
    ]
)

add_paragraph('**Key milestones:**')
add_bullet('**Month 2:** IRB approved, data access established')
add_bullet('**Month 5:** Baseline models trained, initial AUROC benchmarks')
add_bullet('**Month 7:** Full model achieving target AUROC >0.90')
add_bullet('**Month 9:** Primary paper submitted')
add_bullet('**Month 12:** Silent deployment begins')
add_bullet('**Month 18:** Active deployment begins')
add_bullet('**Month 24:** Deployment paper submitted; TFDA registration filed')

doc.add_page_break()

# ============================================================
# 13. RISK MITIGATION
# ============================================================
add_heading('13. Risk Mitigation', 1)

add_table(
    ['Risk', 'Likelihood', 'Impact', 'Mitigation Strategy'],
    [
        ['Data quality issues (missing, inconsistent)', 'High', 'High', 'Extensive EDA in Phase 1; robust missing data handling (GRU-D); sensitivity analyses'],
        ['Model underperformance (<0.90 AUROC)', 'Medium', 'High', 'Multiple architecture options; ensemble methods; feature engineering iteration'],
        ['Poor calibration', 'Medium', 'High', 'Platt scaling or isotonic regression post-hoc; temperature scaling'],
        ['Alert fatigue in deployment', 'High', 'High', 'Tiered alerting; PPV >25% threshold; temporal suppression; iterative threshold tuning'],
        ['TFDA regulatory delays', 'Medium', 'Medium', 'Early pre-submission consultation; leverage NTUH PANCREASaver experience'],
        ['Clinician resistance to AI alerts', 'Medium', 'High', 'Clinical champion program; co-design with ICU nurses/physicians; transparent explanations'],
        ['Distribution shift over time', 'Medium', 'Medium', 'Continuous drift monitoring; automated retraining triggers; quarterly reviews'],
        ['Competition (similar paper published)', 'Low', 'Medium', 'Asian population focus and ECG fusion provide unique angles; move quickly'],
        ['Data leakage in evaluation', 'Low', 'Critical', 'Strict temporal splitting; no patient overlap; code review checklist'],
    ]
)

doc.add_page_break()

# ============================================================
# 14. COMPETITIVE LANDSCAPE
# ============================================================
add_heading('14. Competitive Landscape & Differentiation', 1)

add_heading('14.1 How We Compare', 2)

add_table(
    ['System/Study', 'Data Size', 'Population', 'Modality', 'Full Stay?', 'ECG?', 'Deployed?'],
    [
        ['MIMIC benchmarks', '60K stays', 'US (Boston)', 'Structured', 'No (48h)', 'No', 'No'],
        ['eICU-CRD', '200K stays', 'US (multi-center)', 'Structured', 'No (24h)', 'No', 'No'],
        ['AmsterdamUMCdb', '20K stays', 'European', 'Structured', 'Yes', 'No', 'No'],
        ['Google (Rajkomar 2018)', '216K patients', 'US', 'EHR (FHIR)', 'Yes', 'No', 'No'],
        ['CMUH i.A.M.S.', 'Multi-center', 'Taiwanese', 'Structured', 'Partial', 'No', 'Yes (12 hospitals)'],
        ['Singapore SERA', 'National', 'Singaporean', 'Structured', 'Yes', 'No', 'Yes (national)'],
        ['NTUH (Ours)', '20K patients', 'Taiwanese', 'Structured + ECG', 'Yes', 'YES', 'Planned'],
    ]
)

add_paragraph('**Our unique combination:** Full-stay + ECG waveforms + Asian population + deployment pathway at a top-5 Asian hospital. No existing system combines all four.')

add_heading('14.2 Key Differentiators for Reviewers', 2)

add_bullet('**vs. MIMIC/eICU studies:** Our data is full-stay (not truncated to 48h), includes ECG waveforms, and represents an Asian population \u2014 addressing three major limitations simultaneously')
add_bullet('**vs. CMUH i.A.M.S.:** We include raw ECG waveforms (not just structured data) and focus on a composite deterioration endpoint with specific lead-time analysis')
add_bullet('**vs. Google 2018:** We add ECG waveform modality and validate on Asian population; Google study had no ECG and no non-US validation')
add_bullet('**vs. AmsterdamUMCdb:** Similar scale (~20K) but different population (Asian vs European), and we add ECG waveform data')

doc.add_page_break()

# ============================================================
# 15. REFERENCES
# ============================================================
add_heading('15. References', 1)

refs = [
    'Harutyunyan et al. "Multitask Learning and Benchmarking with Clinical Time Series Data." Scientific Data 6, 96 (2019).',
    'Rajkomar et al. "Scalable and Accurate Deep Learning with Electronic Health Records." npj Digital Medicine 1, 18 (2018).',
    'Che et al. "Recurrent Neural Networks for Multivariate Time Series with Missing Values." Scientific Reports 8, 6085 (2018).',
    'Tipirneni & Reddy. "STraTS: Self-Supervised Transformer for Sparse and Irregularly Sampled Clinical Time-Series." ACM TKDD 16(6), 2022.',
    'Zhang et al. "Raindrop: Graph-Guided Network for Irregularly Sampled Multivariate Time Series." ICLR 2022.',
    'Shukla & Marlin. "Multi-Time Attention Networks for Irregularly Sampled Time Series." ICLR 2021.',
    'De Brouwer et al. "GRU-ODE-Bayes: Continuous Modeling of Sporadically-Observed Time Series." NeurIPS 2019.',
    'Rubanova et al. "Latent ODEs for Irregularly-Sampled Time Series." NeurIPS 2019.',
    'Kidger et al. "Neural Controlled Differential Equations for Irregular Time Series." NeurIPS 2020.',
    'TBAL: "Dynamic Real-Time Risk Prediction Model for ICU In-Hospital Mortality." JMIR Medical Informatics (2025).',
    'RealMIP: "Unlocking the Potential of Real-Time ICU Mortality Prediction with Generative Imputation." npj Digital Medicine (2025).',
    'CRISP: "Causal Relationship Informed Superior Prediction for ICU Mortality." (2024).',
    'XMI-ICU: "Explainable Machine Learning for ICU Mortality in Myocardial Infarction Patients." Scientific Reports (2025).',
    'Dascena InSight. FDA-cleared sepsis prediction system using gradient tree boosting on 6 vital signs.',
    'Wong et al. "External Validation of a Widely Implemented Proprietary Sepsis Prediction Model." JAMA Internal Medicine (2021).',
    'Churpek et al. "Internal and External Validation of a Machine Learning Risk Score for Acute Kidney Injury." JAMA Network Open (2020).',
    'Kendall et al. "Multi-Task Learning Using Uncertainty to Weigh Losses." CVPR 2018.',
    'NTUH AI-Enabled ECG for CAD Prediction. Biomedicines (2022).',
    'NTUH Emergency Triage with Interpretable DL. JMIR Medical Informatics (2024).',
    'CMUH i.A.M.S.: Microsoft News "Inside Taiwan\'s AI Hospital of the Future" (2024).',
    'Singapore BRAIN Platform: Healthcare IT News Asia "Behind Singapore\'s Widespread AI Adoption" (2024).',
    'Stanford Health Care: "Stanford Health Uses AI to Reduce Clinical Deterioration Events." Healthcare IT News (2025).',
    'Meta-analysis: "AI Early Warning Systems and In-Hospital Mortality: Systematic Review." PMC (2025).',
    'Taiwan TFDA SaMD Guidelines. Pacific Bridge Medical (2024); Cisema Updated Guidelines (2025).',
    'Taiwan PDPA Amendments. Jones Day (2025).',
    'Taiwan AI Basic Law. White & Case AI Watch Global Regulatory Tracker (2025).',
    'VGHTPE TvHEWS: "Hemodynamic Instability Prediction in ICU with Time-Varying ML Models." (2024).',
    'eCART at Yale New Haven Health. AgileMD / Sepsis Alliance Webinar Recap (2024).',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.add_run(f'[{i}] ').bold = True
    p.add_run(ref)

# Save
output_path = '/home/user/ECG/NTUH_ICU_AI_Project_Plan.docx'
doc.save(output_path)
print(f"DOCX saved to {output_path}")
